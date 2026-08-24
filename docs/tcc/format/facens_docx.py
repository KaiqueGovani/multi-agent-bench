#!/usr/bin/env python3
"""Formata e valida o DOCX do TCC segundo o manual técnico da Facens.

Fluxo determinístico em duas passagens:

1. ``prepare`` normaliza estilos, seções, numeração, referências e cria um
   sumário estático provisório com líderes de tabulação automáticos.
2. O DOCX é renderizado em PDF.
3. ``finalize-toc`` lê as páginas efetivas do PDF e grava os números corretos
   no sumário, preservando links internos para os títulos.
4. ``audit`` verifica estrutura, referências e, se fornecido, paginação.

O sumário estático é intencional: evita a dependência de atualização manual de
campos do Word/LibreOffice em execuções headless do Ralph loop.
"""

from __future__ import annotations

import argparse
import copy
import re
import sys
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HEADING_STYLES = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}
TOC_STYLES = {1: "Facens TOC 1", 2: "Facens TOC 2", 3: "Facens TOC 3"}
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
LEFT_MARGIN_CM = 3.0
RIGHT_MARGIN_CM = 2.0
TOP_MARGIN_CM = 3.0
BOTTOM_MARGIN_CM = 2.0
TOC_RIGHT_TAB_CM = PAGE_WIDTH_CM - LEFT_MARGIN_CM - RIGHT_MARGIN_CM
TOC_INDENTS_CM = {1: 0.0, 2: 0.75, 3: 1.5}
ET_AL_RE = re.compile(r"\bet\s+al\.", re.IGNORECASE)
MANUAL_NUMBER_RE = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s+")
FOREIGN_TERM_RE = re.compile(
    r"\b(?:"
    r"large\s+language\s+models?|multi-agent\s+benchmark|customer\s+care|"
    r"service\s+operations|server-sent\s+events|llm-as-judge|scoping\s+review|"
    r"recall@3|swarm\s+intelligence|chatbots?|benchmarks?|pipelines?|workflows?|"
    r"swarms?|handoffs?|runs?|seeds?|hooks?|backends?|frontends?|stacks?|"
    r"dashboards?|workspaces?|runtimes?|loops?|timeouts?|endpoints?|baselines?|"
    r"centralized_orchestration|structured_workflow|decentralized_swarm|"
    r"faq_lookup|stock_lookup|attachment_intake|live|mock|recall|chat|software"
    r")\b",
    re.IGNORECASE,
)

# A NBR 6023:2018 permite diferentes recursos tipográficos para o elemento
# título, desde que o recurso escolhido seja uniforme. O manual Facens usa
# negrito nos exemplos; em artigos, o destaque recai no título do periódico.
REFERENCE_EMPHASIS = {
    "ADA HEALTH.": "Health. Powered by Ada",
    "ALVAREZ,": "Exploratory Research in Clinical and Social Pharmacy",
    "AMAZON.": "Pharmacy AI Assistant FAQs",
    "AMAZON SCIENCE.": "The life of a prescription at Amazon Pharmacy",
    "ANTHROPIC.": "Building effective AI agents",
    "BALAPRAKASH": "International Journal of High Performance Computing Applications",
    "BUESING": "Where is customer care in 2024?",
    "FIP.": "An artificial intelligence toolkit for pharmacy: an introduction and resource guide for pharmacists",
    "HAMMOND": "Multi-Agent Risks from Advanced AI",
    "HATZIMANOLIS": "Research in Social and Administrative Pharmacy",
    "INFERMEDICA.": "Symptom Checker",
    "INTERCOM.": "Fin AI Agent explained",
    "JAGATAP": "Proceedings of the 2025 Conference of the Nations of the Americas Chapter of the Association for Computational Linguistics: Human Language Technologies",
    "KIM": "Nature Machine Intelligence",
    "LI": "Vicinagearth",
    "MCKINSEY & COMPANY.": "The state of AI in early 2024: gen AI adoption spikes and starts to generate value",
    "PGEU.": "Position paper on artificial intelligence",
    "PICCIALLI": "Expert Systems with Applications",
    "SALESFORCE.": "AI Customer Service Agents",
    "STANFORD HAI.": "AI Index Report 2025",
    "WHO EUROPE.": "Advancing the role of pharmacists to meet changing patient and health system needs",
    "YU": "2025 8th International Conference on Artificial Intelligence and Big Data (ICAIBD)",
    "ZENDESK.": "About AI agents",
}


@dataclass(frozen=True)
class TocEntry:
    level: int
    title: str
    display: str
    bookmark: str
    page: int | None = None


def normalized(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u00a0", " ").strip()).upper()


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def iter_all_paragraphs(doc: Document):
    """Percorre corpo e tabelas sem limitar a regra tipográfica ao texto corrido."""
    yield from doc.paragraphs
    for table in doc.tables:
        yield from iter_table_paragraphs(table)


def set_fonts(element, family: str = "Arial") -> None:
    r_pr = element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{name}"), family)


def set_style_font(style, size: int, *, bold: bool = False, caps: bool = False) -> None:
    style.font.name = "Arial"
    set_fonts(style._element)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.all_caps = caps
    style.font.color.rgb = RGBColor(0, 0, 0)
    remove_theme_color(style._element.get_or_add_rPr())


def remove_theme_color(r_pr) -> None:
    """Força preto absoluto; Word prioriza themeColor sobre w:val em alguns casos."""
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.set(qn("w:val"), "000000")
    for attribute in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(f"w:{attribute}"), None)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.left_indent = Cm(0)
    normal.paragraph_format.right_indent = Cm(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    specifications = {
        "Heading 1": (14, True, True, True),
        "Heading 2": (12, True, True, False),
        "Heading 3": (12, False, True, False),
    }
    for name, (size, bold, caps, new_page) in specifications.items():
        style = styles[name]
        set_style_font(style, size, bold=bold, caps=caps)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
        fmt.left_indent = Cm(0)
        fmt.right_indent = Cm(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(18)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.keep_with_next = True
        fmt.keep_together = True
        fmt.page_break_before = new_page

    for style_name in ["Front Matter Heading", "Reference Entry", *TOC_STYLES.values()]:
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    front = styles["Front Matter Heading"]
    set_style_font(front, 14, bold=True, caps=True)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.first_line_indent = Cm(0)
    front.paragraph_format.space_before = Pt(0)
    front.paragraph_format.space_after = Pt(18)
    front.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    front.paragraph_format.keep_with_next = True

    reference = styles["Reference Entry"]
    set_style_font(reference, 12)
    reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reference.paragraph_format.first_line_indent = Cm(0)
    reference.paragraph_format.left_indent = Cm(0)
    reference.paragraph_format.right_indent = Cm(0)
    reference.paragraph_format.space_before = Pt(0)
    reference.paragraph_format.space_after = Pt(12)
    reference.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    reference.paragraph_format.keep_together = True

    for level, style_name in TOC_STYLES.items():
        style = styles[style_name]
        set_style_font(style, 12, bold=(level == 1), caps=False)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
        fmt.left_indent = Cm(TOC_INDENTS_CM[level])
        fmt.right_indent = Cm(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.keep_together = True
        fmt.tab_stops.clear_all()
        fmt.tab_stops.add_tab_stop(
            # Todos os números de página terminam na mesma margem direita.
            # O recuo altera apenas o início do título, nunca o fim do pontilhado.
            Cm(TOC_RIGHT_TAB_CM),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )


def set_run_typography(run, size: int, *, bold: bool | None = None, caps: bool | None = None) -> None:
    """Remove a dependência de fontes diretas herdadas do modelo antigo."""
    run.font.name = "Arial"
    set_fonts(run._element)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    remove_theme_color(run._r.get_or_add_rPr())
    if bold is not None:
        run.font.bold = bold
    if caps is not None:
        run.font.all_caps = caps


def enforce_document_typography(doc: Document) -> None:
    """Aplica a hierarquia Facens ao texto visível, inclusive em tabelas."""
    body_started = False
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if style_name == "Heading 1":
            body_started = True

        if style_name == "Heading 1":
            for run in paragraph.runs:
                set_run_typography(run, 14, bold=True, caps=True)
        elif style_name == "Heading 2":
            for run in paragraph.runs:
                set_run_typography(run, 12, bold=True, caps=True)
        elif style_name == "Heading 3":
            for run in paragraph.runs:
                set_run_typography(run, 12, bold=False, caps=True)
        elif style_name == "Front Matter Heading":
            for run in paragraph.runs:
                set_run_typography(run, 14, bold=True, caps=True)
        elif style_name in TOC_STYLES.values():
            level = next(level for level, name in TOC_STYLES.items() if name == style_name)
            for run in paragraph.runs:
                set_run_typography(run, 12, bold=(level == 1), caps=False)
        elif style_name == "Reference Entry" or body_started:
            for run in paragraph.runs:
                set_run_typography(run, 12)

    # As tabelas pertencem ao corpo acadêmico nesta versão do documento.
    for table in doc.tables:
        for paragraph in iter_table_paragraphs(table):
            for run in paragraph.runs:
                set_run_typography(run, 12)


def configure_heading_spacing(doc: Document) -> None:
    """Mantém uma linha e meia antes/depois dos títulos sem duplicar intervalos."""
    previous_was_heading = False
    for paragraph in doc.paragraphs:
        level = HEADING_STYLES.get(paragraph.style.name)
        if level is None:
            previous_was_heading = False
            continue
        fmt = paragraph.paragraph_format
        # Títulos de nível 1 já iniciam nova página. Nos demais, 18 pt equivale
        # ao intervalo visual de uma linha e meia em Arial 12.
        fmt.space_before = Pt(0 if level == 1 or previous_was_heading else 18)
        fmt.space_after = Pt(18)
        previous_was_heading = True


def set_xml_run_property(run_element, tag: str) -> None:
    r_pr = run_element.get_or_add_rPr()
    prop = r_pr.find(qn(f"w:{tag}"))
    if prop is None:
        prop = OxmlElement(f"w:{tag}")
        r_pr.append(prop)
    prop.set(qn("w:val"), "1")


def format_regex_in_paragraph(paragraph, pattern: re.Pattern, *, italic=False, bold=False) -> None:
    """Formata correspondências inclusive dentro de hyperlinks do sumário."""
    for run_element in list(paragraph._p.findall(".//" + qn("w:r"))):
        text_nodes = run_element.findall(".//" + qn("w:t"))
        text = "".join(node.text or "" for node in text_nodes)
        matches = list(pattern.finditer(text))
        if not matches:
            continue
        parent = run_element.getparent()
        position = parent.index(run_element)
        spans = [(match.start(), match.end()) for match in matches]
        boundaries = sorted({0, len(text), *(value for span in spans for value in span)})
        rebuilt = []
        for start, end in zip(boundaries, boundaries[1:]):
            if start == end:
                continue
            new_run = copy.deepcopy(run_element)
            for node in list(new_run.findall(".//" + qn("w:t"))):
                node.getparent().remove(node)
            segment = text[start:end]
            is_match = any(
                start >= match_start and end <= match_end
                for match_start, match_end in spans
            )
            if is_match:
                if italic:
                    set_xml_run_property(new_run, "i")
                    set_xml_run_property(new_run, "iCs")
                if bold:
                    set_xml_run_property(new_run, "b")
                    set_xml_run_property(new_run, "bCs")
            text_node = OxmlElement("w:t")
            text_node.text = segment
            if segment[:1].isspace() or segment[-1:].isspace():
                text_node.set(qn("xml:space"), "preserve")
            new_run.append(text_node)
            rebuilt.append(new_run)
        parent.remove(run_element)
        for offset, new_run in enumerate(rebuilt):
            parent.insert(position + offset, new_run)


def italicize_foreign_terms(doc: Document) -> None:
    """Aplica itálico a estrangeirismos isolados, sem reformatar o ABSTRACT."""
    body_started = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            body_started = True
        if normalized(paragraph.text) == "REFERÊNCIAS":
            body_started = False
        if body_started or paragraph.style.name in TOC_STYLES.values():
            format_regex_in_paragraph(paragraph, FOREIGN_TERM_RE, italic=True)
    for table in doc.tables:
        for paragraph in iter_table_paragraphs(table):
            format_regex_in_paragraph(paragraph, FOREIGN_TERM_RE, italic=True)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)
        section.top_margin = Cm(TOP_MARGIN_CM)
        section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    if len(doc.sections) < 3:
        raise ValueError("O documento precisa preservar as seções de capa, pré-textuais e texto.")

    # A capa não é contada. A folha de rosto (segunda seção) inicia a contagem em 1.
    prelim = doc.sections[1]._sectPr
    pg_num = prelim.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        prelim.append(pg_num)
    pg_num.set(qn("w:start"), "1")

    # A seção textual continua a contagem; não usa um valor fixo, pois o sumário
    # pode crescer durante o Ralph loop.
    body = doc.sections[2]._sectPr
    body_pg_num = body.find(qn("w:pgNumType"))
    if body_pg_num is not None:
        body_pg_num.attrib.pop(qn("w:start"), None)

    for index, section in enumerate(doc.sections):
        section.footer.is_linked_to_previous = False
        if index < 2:
            for paragraph in section.footer.paragraphs:
                paragraph.clear()
        else:
            footer = section.footer
            while len(footer.paragraphs) > 1:
                footer._element.remove(footer.paragraphs[-1]._p)
            paragraph = footer.paragraphs[0]
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            run.font.name = "Arial"
            set_fonts(run._element)
            run.font.size = Pt(10)
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = " PAGE "
            separate = OxmlElement("w:fldChar")
            separate.set(qn("w:fldCharType"), "separate")
            cached = OxmlElement("w:t")
            cached.text = "1"
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            run._r.extend([begin, instruction, separate, cached, end])


def numbering_definition_is_facens(numbering, num_id: str) -> bool:
    num = numbering.find(f"./{qn('w:num')}[@{qn('w:numId')}='{num_id}']")
    if num is None:
        return False
    abstract_ref = num.find(qn("w:abstractNumId"))
    if abstract_ref is None:
        return False
    abstract_id = abstract_ref.get(qn("w:val"))
    abstract = numbering.find(
        f"./{qn('w:abstractNum')}[@{qn('w:abstractNumId')}='{abstract_id}']"
    )
    if abstract is None:
        return False
    expected_texts = ["%1", "%1.%2", "%1.%2.%3"]
    for level in range(3):
        lvl = abstract.find(f"./{qn('w:lvl')}[@{qn('w:ilvl')}='{level}']")
        if lvl is None:
            return False
        num_fmt = lvl.find(qn("w:numFmt"))
        lvl_text = lvl.find(qn("w:lvlText"))
        p_style = lvl.find(qn("w:pStyle"))
        if (
            num_fmt is None
            or num_fmt.get(qn("w:val")) != "decimal"
            or lvl_text is None
            or lvl_text.get(qn("w:val")) != expected_texts[level]
            or p_style is None
            or p_style.get(qn("w:val")) != f"Heading{level + 1}"
        ):
            return False
    return True


def create_numbering_definition(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), "FAC30001")
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), "FAC30001")
    abstract.extend([nsid, multi, template])
    for level, level_text in enumerate(["%1", "%1.%2", "%1.%2.%3"]):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "decimal")
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), f"Heading{level + 1}")
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), level_text)
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "0")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:hanging"), "0")
        p_pr.extend([tabs, ind])
        # A ordem segue CT_Lvl. O Word pode descartar vínculos de estilo que
        # apareçam fora da sequência OOXML, fazendo a lista parecer um bullet
        # ou deixando os níveis inferiores sem numeração na interface.
        lvl.extend([start, fmt, p_style, suffix, text, justification, p_pr])
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def make_heading_num_pr(level: int, num_id: int):
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level - 1))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    return num_pr


def replace_num_pr(p_pr, num_pr) -> None:
    old = p_pr.find(qn("w:numPr"))
    if old is not None:
        p_pr.remove(old)
    # Usa o ponto de inserção definido pelo schema, evitando numPr fora de
    # ordem tanto em estilos quanto em parágrafos.
    p_pr._insert_numPr(num_pr)


def bind_heading_styles_to_numbering(doc: Document, num_id: int) -> None:
    """Vincula a lista decimal aos estilos, não só aos parágrafos existentes."""
    for style_name, level in HEADING_STYLES.items():
        style_p_pr = doc.styles[style_name]._element.get_or_add_pPr()
        replace_num_pr(style_p_pr, make_heading_num_pr(level, num_id))


def configure_heading_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    num_id = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name not in HEADING_STYLES:
            continue
        p_pr = paragraph._p.pPr
        if p_pr is None or p_pr.numPr is None or p_pr.numPr.numId is None:
            continue
        candidate = str(p_pr.numPr.numId.val)
        if numbering_definition_is_facens(numbering, candidate):
            num_id = int(candidate)
            break
    if num_id is None:
        num_id = create_numbering_definition(doc)

    bind_heading_styles_to_numbering(doc, num_id)

    for paragraph in doc.paragraphs:
        level = HEADING_STYLES.get(paragraph.style.name)
        if level is None:
            continue
        title = MANUAL_NUMBER_RE.sub("", paragraph.text, count=1)
        # A caixa alta é gravada no texto, não apenas simulada pelo estilo. Isso
        # mantém a mesma grafia no corpo, no sumário e em outros leitores DOCX.
        title = re.sub(r"\s+", " ", title.replace("\u00a0", " ").strip()).upper()
        if paragraph.text != title:
            paragraph.text = title
        p_pr = paragraph._p.get_or_add_pPr()
        replace_num_pr(p_pr, make_heading_num_pr(level, num_id))
    return num_id


def replace_reference_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)
    paragraph.style = "Reference Entry"


def italicize_et_al(paragraph) -> None:
    for run in list(paragraph.runs):
        if not ET_AL_RE.search(run.text):
            continue
        parent = run._r.getparent()
        position = parent.index(run._r)
        pieces = ET_AL_RE.split(run.text)
        matches = list(ET_AL_RE.finditer(run.text))
        rebuilt = []
        for index, piece in enumerate(pieces):
            if piece:
                new_run = copy.deepcopy(run._r)
                for text_node in new_run.findall(".//" + qn("w:t")):
                    text_node.getparent().remove(text_node)
                text_node = OxmlElement("w:t")
                text_node.text = piece
                if piece[:1].isspace() or piece[-1:].isspace():
                    text_node.set(qn("xml:space"), "preserve")
                new_run.append(text_node)
                rebuilt.append(new_run)
            if index < len(matches):
                new_run = copy.deepcopy(run._r)
                for text_node in new_run.findall(".//" + qn("w:t")):
                    text_node.getparent().remove(text_node)
                r_pr = new_run.get_or_add_rPr()
                italic = r_pr.find(qn("w:i"))
                if italic is None:
                    italic = OxmlElement("w:i")
                    r_pr.append(italic)
                italic.set(qn("w:val"), "1")
                italic_cs = r_pr.find(qn("w:iCs"))
                if italic_cs is None:
                    italic_cs = OxmlElement("w:iCs")
                    r_pr.append(italic_cs)
                italic_cs.set(qn("w:val"), "1")
                text_node = OxmlElement("w:t")
                text_node.text = matches[index].group(0)
                new_run.append(text_node)
                rebuilt.append(new_run)
        parent.remove(run._r)
        for offset, rebuilt_run in enumerate(rebuilt):
            parent.insert(position + offset, rebuilt_run)


def configure_references(doc: Document) -> None:
    heading = next(
        (p for p in doc.paragraphs if normalized(p.text) == "REFERÊNCIAS"), None
    )
    if heading is None:
        raise ValueError("Título REFERÊNCIAS não encontrado.")
    heading.style = "Front Matter Heading"
    heading.paragraph_format.page_break_before = True

    corrections = {
        "BALAPRAKASH, PRASANNA": (
            "BALAPRAKASH, Prasanna et al. SWARM: reimagining scientific workflow "
            "management systems in a distributed world. International Journal of High "
            "Performance Computing Applications, [S. l.], v. 39, n. 5, p. 692-712, "
            "2025. DOI: 10.1177/10943420251339317. Disponível em: "
            "https://doi.org/10.1177/10943420251339317. Acesso em: 30 mar. 2026."
        ),
        "HATZIMANOLIS, ANDREW": (
            "HATZIMANOLIS, Andrew et al. Applications of artificial intelligence in "
            "current pharmacy practice: a scoping review. Research in Social and "
            "Administrative Pharmacy, [S. l.], v. 21, n. 3, p. 134-141, 2025. DOI: "
            "10.1016/j.sapharm.2024.12.007. Disponível em: "
            "https://doi.org/10.1016/j.sapharm.2024.12.007. Acesso em: 30 mar. 2026."
        ),
    }
    entries = []
    seen = False
    for paragraph in doc.paragraphs:
        if paragraph._p is heading._p:
            seen = True
            continue
        if not seen or not paragraph.text.strip():
            continue
        upper = normalized(paragraph.text)
        for prefix, replacement in corrections.items():
            if upper.startswith(prefix):
                replace_reference_text(paragraph, replacement)
                break
        # Reconstrói a entrada para remover formatações residuais do modelo e
        # reaplicar de modo determinístico o destaque bibliográfico Facens.
        reference_text = paragraph.text
        replace_reference_text(paragraph, reference_text)
        paragraph.style = "Reference Entry"
        paragraph.paragraph_format.page_break_before = False
        upper = normalized(paragraph.text)
        emphasis = next(
            (title for prefix, title in REFERENCE_EMPHASIS.items() if upper.startswith(prefix)),
            None,
        )
        if emphasis and emphasis in paragraph.text:
            format_regex_in_paragraph(
                paragraph, re.compile(re.escape(emphasis)), bold=True
            )
        italicize_et_al(paragraph)
        entries.append(paragraph)

    entries.sort(key=lambda p: normalized(p.text).split(".", 1)[0])
    parent = heading._p.getparent()
    for paragraph in entries:
        parent.remove(paragraph._p)
    anchor = heading._p
    for paragraph in entries:
        anchor.addnext(paragraph._p)
        anchor = paragraph._p

    # A regra vale também para citações no corpo.
    for paragraph in iter_all_paragraphs(doc):
        italicize_et_al(paragraph)


def remove_own_bookmarks(paragraph) -> None:
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for element in list(paragraph._p.findall(".//" + qn(tag))):
            if tag == "w:bookmarkStart" and not element.get(qn("w:name"), "").startswith(
                "_TCC_H_"
            ):
                continue
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    remove_own_bookmarks(paragraph)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p_pr = paragraph._p.pPr
    insert_at = 1 if p_pr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def collect_toc_entries(doc: Document) -> list[TocEntry]:
    entries = []
    counters = [0, 0, 0]
    started = False
    bookmark_id = 1000
    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        if style == "Heading 1":
            started = True
        if not started:
            continue
        level = HEADING_STYLES.get(style)
        if level:
            counters[level - 1] += 1
            for index in range(level, 3):
                counters[index] = 0
            number = ".".join(str(value) for value in counters[:level])
            title = re.sub(r"\s+", " ", paragraph.text.replace("\u00a0", " ").strip())
            bookmark = f"_TCC_H_{bookmark_id:04d}"
            add_bookmark(paragraph, bookmark, bookmark_id)
            entries.append(TocEntry(level, title, f"{number} {title}", bookmark))
            bookmark_id += 1
            continue
        if normalized(paragraph.text) == "REFERÊNCIAS":
            bookmark = f"_TCC_H_{bookmark_id:04d}"
            add_bookmark(paragraph, bookmark, bookmark_id)
            entries.append(TocEntry(1, "REFERÊNCIAS", "REFERÊNCIAS", bookmark))
            break
    return entries


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for name in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{name}"), "Arial")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.extend([fonts, color, underline, size])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set(qn("xml:space"), "preserve")
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def toc_boundary(doc: Document, toc_heading) -> object:
    cursor = toc_heading._p.getnext()
    while cursor is not None:
        p_pr = cursor.find(qn("w:pPr"))
        if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
            return cursor
        cursor = cursor.getnext()
    raise ValueError("Quebra de seção após o SUMÁRIO não encontrada.")


def rebuild_toc(doc: Document, entries: list[TocEntry]) -> None:
    toc_heading = next(
        (p for p in doc.paragraphs if normalized(p.text) == "SUMÁRIO"), None
    )
    if toc_heading is None:
        raise ValueError("Título SUMÁRIO não encontrado.")
    toc_heading.style = "Front Matter Heading"
    toc_heading.paragraph_format.page_break_before = True
    boundary = toc_boundary(doc, toc_heading)
    cursor = toc_heading._p.getnext()
    while cursor is not None and cursor is not boundary:
        following = cursor.getnext()
        cursor.getparent().remove(cursor)
        cursor = following

    for entry in entries:
        paragraph = doc.add_paragraph(style=TOC_STYLES[entry.level])
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(TOC_RIGHT_TAB_CM),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        add_internal_hyperlink(paragraph, entry.display, entry.bookmark)
        page_text = str(entry.page) if entry.page is not None else "—"
        separator = paragraph.add_run("\t")
        set_run_typography(separator, 12, bold=(entry.level == 1), caps=False)
        add_internal_hyperlink(paragraph, page_text, entry.bookmark)
        boundary.addprevious(paragraph._p)


def page_map_from_pdf(pdf_path: Path, entries: list[TocEntry]) -> dict[str, int]:
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            pages.append(normalized(page.extract_text() or ""))
    page_map = {}
    for entry in entries:
        needle = normalized(entry.display)
        matches = [index + 1 for index, text in enumerate(pages) if needle in text]
        if not matches:
            raise ValueError(f"Título não localizado no PDF: {entry.display}")
        physical_page = max(matches)
        # O manual Facens exclui a capa da contagem.
        page_map[entry.bookmark] = physical_page - 1
    return page_map


def prepare(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    configure_styles(doc)
    configure_sections(doc)
    configure_heading_numbering(doc)
    configure_heading_spacing(doc)
    configure_references(doc)
    entries = collect_toc_entries(doc)
    rebuild_toc(doc, entries)
    enforce_document_typography(doc)
    italicize_foreign_terms(doc)
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def finalize_toc(input_path: Path, pdf_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    entries = collect_toc_entries(doc)
    page_map = page_map_from_pdf(pdf_path, entries)
    final_entries = [
        TocEntry(entry.level, entry.title, entry.display, entry.bookmark, page_map[entry.bookmark])
        for entry in entries
    ]
    rebuild_toc(doc, final_entries)
    enforce_document_typography(doc)
    configure_heading_spacing(doc)
    italicize_foreign_terms(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def run_is_italic(run) -> bool:
    if run.italic is not None:
        return bool(run.italic)
    r_pr = run._r.rPr
    if r_pr is None or r_pr.find(qn("w:i")) is None:
        return False
    return r_pr.find(qn("w:i")).get(qn("w:val"), "1") not in {"0", "false", "off"}


def xml_run_has_property(run_element, tag: str) -> bool:
    r_pr = run_element.find(qn("w:rPr"))
    prop = r_pr.find(qn(f"w:{tag}")) if r_pr is not None else None
    return prop is not None and prop.get(qn("w:val"), "1") not in {
        "0",
        "false",
        "off",
    }


def audit(input_path: Path, pdf_path: Path | None = None) -> list[str]:
    doc = Document(input_path)
    errors = []
    tolerance = 0.03
    for index, section in enumerate(doc.sections, start=1):
        expected = [21.0, 29.7, 3.0, 2.0, 3.0, 2.0]
        actual = [
            section.page_width.cm,
            section.page_height.cm,
            section.left_margin.cm,
            section.right_margin.cm,
            section.top_margin.cm,
            section.bottom_margin.cm,
        ]
        if any(abs(a - e) > tolerance for a, e in zip(actual, expected)):
            errors.append(f"Seção {index}: geometria fora do padrão Facens: {actual}")

    headings = [p for p in doc.paragraphs if p.style.name in HEADING_STYLES]
    if not headings:
        errors.append("Nenhum título estruturado foi encontrado.")
    num_ids = set()
    expected_heading_styles = {
        "Heading 1": (14, True, True),
        "Heading 2": (12, True, True),
        "Heading 3": (12, False, True),
    }
    for style_name, (size, bold, caps) in expected_heading_styles.items():
        style = doc.styles[style_name]
        if style.font.name != "Arial" or style.font.size != Pt(size):
            errors.append(f"{style_name}: fonte deve ser Arial {size}.")
        if bool(style.font.bold) != bold:
            errors.append(f"{style_name}: negrito incorreto.")
        if bool(style.font.all_caps) != caps:
            errors.append(f"{style_name}: caixa alta incorreta.")
    for paragraph in headings:
        if paragraph.text != paragraph.text.upper():
            errors.append(f"Título fora de caixa alta: {paragraph.text}")
        if MANUAL_NUMBER_RE.match(paragraph.text):
            errors.append(f"Numeração manual em título: {paragraph.text}")
        expected_level = HEADING_STYLES[paragraph.style.name] - 1
        p_pr = paragraph._p.pPr
        if p_pr is None or p_pr.numPr is None or p_pr.numPr.numId is None:
            errors.append(f"Título sem numeração multinível: {paragraph.text}")
        else:
            num_ids.add(str(p_pr.numPr.numId.val))
            if p_pr.numPr.ilvl is None or int(p_pr.numPr.ilvl.val) != expected_level:
                errors.append(
                    f"Título no nível numérico incorreto: {paragraph.text} "
                    f"(esperado {expected_level})"
                )
    if len(num_ids) != 1:
        errors.append(f"Títulos usam mais de uma lista multinível: {sorted(num_ids)}")
    else:
        heading_num_id = next(iter(num_ids))
        numbering = doc.part.numbering_part.element
        if not numbering_definition_is_facens(numbering, heading_num_id):
            errors.append(
                "Lista dos títulos não possui três níveis decimais vinculados "
                "a Heading 1, Heading 2 e Heading 3."
            )
        for style_name, level in HEADING_STYLES.items():
            style_p_pr = doc.styles[style_name]._element.pPr
            style_num_pr = (
                style_p_pr.find(qn("w:numPr")) if style_p_pr is not None else None
            )
            style_ilvl = (
                style_num_pr.find(qn("w:ilvl")) if style_num_pr is not None else None
            )
            style_num_id = (
                style_num_pr.find(qn("w:numId")) if style_num_pr is not None else None
            )
            if (
                style_ilvl is None
                or style_num_id is None
                or style_ilvl.get(qn("w:val")) != str(level - 1)
                or style_num_id.get(qn("w:val")) != heading_num_id
            ):
                errors.append(
                    f"{style_name}: estilo não vinculado ao nível decimal {level} "
                    "da lista única dos títulos."
                )

    toc_paragraphs = [p for p in doc.paragraphs if p.style.name in TOC_STYLES.values()]
    entries = collect_toc_entries(doc)
    if len(toc_paragraphs) != len(entries):
        errors.append(
            f"Sumário incompleto: {len(toc_paragraphs)} entradas para {len(entries)} títulos."
        )
    if any("—" in p.text or "Atualize o sumário" in p.text for p in toc_paragraphs):
        errors.append("Sumário ainda contém marcador provisório.")
    for paragraph in toc_paragraphs:
        if re.search(r"\.{4,}", paragraph.text):
            errors.append(f"Entrada do sumário simula líder com pontos literais: {paragraph.text}")
        p_pr = paragraph._p.pPr
        tabs = p_pr.find(qn("w:tabs")) if p_pr is not None else None
        right_dot_tabs = [] if tabs is None else [
            tab for tab in tabs.findall(qn("w:tab"))
            if tab.get(qn("w:val")) == "right"
            and tab.get(qn("w:leader")) == "dot"
        ]
        if len(right_dot_tabs) != 1:
            errors.append(
                f"Entrada do sumário sem tabulador direito com líder automático: {paragraph.text}"
            )
        content_tabs = paragraph._p.findall(
            ".//" + qn("w:r") + "/" + qn("w:tab")
        )
        if len(content_tabs) != 1:
            errors.append(
                f"Entrada do sumário sem separador de tabulação único: {paragraph.text}"
            )

    for paragraph in headings:
        for run in paragraph.runs:
            r_pr = run._r.rPr
            color = r_pr.find(qn("w:color")) if r_pr is not None else None
            if color is None or color.get(qn("w:val")) != "000000":
                errors.append(f"Título sem preto absoluto: {paragraph.text}")
                break
            if any(color.get(qn(f"w:{name}")) for name in ("themeColor", "themeTint", "themeShade")):
                errors.append(f"Título ainda depende de cor de tema: {paragraph.text}")
                break

    previous_was_heading = False
    for paragraph in doc.paragraphs:
        level = HEADING_STYLES.get(paragraph.style.name)
        if level is None:
            previous_was_heading = False
            continue
        before = paragraph.paragraph_format.space_before.pt
        after = paragraph.paragraph_format.space_after.pt
        expected_before = 0 if level == 1 or previous_was_heading else 18
        if abs(before - expected_before) > 0.1 or abs(after - 18) > 0.1:
            errors.append(f"Espaçamento incorreto no título: {paragraph.text}")
        previous_was_heading = True

    reference_heading = next(
        (p for p in doc.paragraphs if normalized(p.text) == "REFERÊNCIAS"), None
    )
    if reference_heading is None:
        errors.append("Seção REFERÊNCIAS ausente.")
    elif not reference_heading.paragraph_format.page_break_before:
        errors.append("REFERÊNCIAS não inicia em nova página.")

    reference_entries = []
    in_references = False
    for paragraph in doc.paragraphs:
        if reference_heading is not None and paragraph._p is reference_heading._p:
            in_references = True
            continue
        if in_references and paragraph.text.strip():
            reference_entries.append(paragraph)
            if paragraph.style.name != "Reference Entry":
                errors.append(f"Referência com estilo incorreto: {paragraph.text[:80]}")
    keys = [normalized(p.text).split(".", 1)[0] for p in reference_entries]
    if keys != sorted(keys):
        errors.append("Referências não estão em ordem alfabética.")
    for paragraph in reference_entries:
        upper = normalized(paragraph.text)
        emphasis = next(
            (title for prefix, title in REFERENCE_EMPHASIS.items() if upper.startswith(prefix)),
            None,
        )
        if not emphasis:
            errors.append(f"Referência sem regra de destaque: {paragraph.text[:80]}")
            continue
        matched = False
        for run_element in paragraph._p.findall(".//" + qn("w:r")):
            run_text = "".join(
                node.text or ""
                for node in run_element.findall(".//" + qn("w:t"))
            )
            if emphasis in run_text:
                matched = True
                if not xml_run_has_property(run_element, "b"):
                    errors.append(f"Título bibliográfico sem negrito: {emphasis}")
                break
        if not matched:
            errors.append(f"Destaque bibliográfico não localizado: {emphasis}")

    for p_index, paragraph in enumerate(iter_all_paragraphs(doc)):
        if not ET_AL_RE.search(paragraph.text):
            continue
        for run in paragraph.runs:
            if ET_AL_RE.search(run.text) and not run_is_italic(run):
                errors.append(f"'et al.' sem itálico no parágrafo {p_index}: {paragraph.text[:90]}")

    body_started = False
    for p_index, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "Heading 1":
            body_started = True
        if normalized(paragraph.text) == "REFERÊNCIAS":
            body_started = False
        if not (body_started or paragraph.style.name in TOC_STYLES.values()):
            continue
        for run_element in paragraph._p.findall(".//" + qn("w:r")):
            run_text = "".join(
                node.text or ""
                for node in run_element.findall(".//" + qn("w:t"))
            )
            if FOREIGN_TERM_RE.search(run_text) and not xml_run_has_property(run_element, "i"):
                errors.append(f"Estrangeirismo sem itálico no parágrafo {p_index}: {run_text}")
    for table in doc.tables:
        for paragraph in iter_table_paragraphs(table):
            for run_element in paragraph._p.findall(".//" + qn("w:r")):
                run_text = "".join(
                    node.text or ""
                    for node in run_element.findall(".//" + qn("w:t"))
                )
                if FOREIGN_TERM_RE.search(run_text) and not xml_run_has_property(run_element, "i"):
                    errors.append(f"Estrangeirismo sem itálico em tabela: {run_text}")

    if pdf_path is not None:
        try:
            page_map = page_map_from_pdf(pdf_path, entries)
        except ValueError as exc:
            errors.append(str(exc))
        else:
            toc_pages = {}
            for paragraph, entry in zip(toc_paragraphs, entries):
                match = re.search(r"(\d+)\s*$", paragraph.text)
                if match:
                    toc_pages[entry.bookmark] = int(match.group(1))
            for entry in entries:
                if toc_pages.get(entry.bookmark) != page_map[entry.bookmark]:
                    errors.append(
                        f"Página divergente no sumário: {entry.display} "
                        f"({toc_pages.get(entry.bookmark)} != {page_map[entry.bookmark]})"
                    )
    return errors


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)

    prepare_parser = subparsers.add_parser("prepare")
    prepare_parser.add_argument("input", type=Path)
    prepare_parser.add_argument("output", type=Path)

    toc_parser = subparsers.add_parser("finalize-toc")
    toc_parser.add_argument("input", type=Path)
    toc_parser.add_argument("pdf", type=Path)
    toc_parser.add_argument("output", type=Path)

    audit_parser = subparsers.add_parser("audit")
    audit_parser.add_argument("input", type=Path)
    audit_parser.add_argument("--pdf", type=Path)

    args = parser.parse_args()
    if args.command == "prepare":
        prepare(args.input, args.output)
        print(args.output)
        return 0
    if args.command == "finalize-toc":
        finalize_toc(args.input, args.pdf, args.output)
        print(args.output)
        return 0
    errors = audit(args.input, args.pdf)
    if errors:
        for error in errors:
            print(f"ERRO: {error}")
        return 1
    print("OK: DOCX em conformidade estrutural com o perfil Facens.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
