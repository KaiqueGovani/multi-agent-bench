#!/usr/bin/env python3
"""Sincroniza um capítulo Markdown com o DOCX e preserva figuras vetoriais.

O formato aceito é o usado em ``docs/tcc/manuscript``: títulos Markdown,
parágrafos, listas, imagens e comentários ``FIGURE``. As figuras são inseridas
como PNG inline para compatibilidade e recebem o SVG original no ``svgBlip`` do
OOXML; Word usa o vetor e renderizadores antigos continuam com o PNG.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.shared import Cm
from lxml import etree


FIGURE_COMMENT_RE = re.compile(
    r"<!--\s*FIGURE:\s*(?P<id>[^|]+)\|\s*caption:\s*(?P<caption>.*?)\s*-->",
    re.IGNORECASE,
)
IMAGE_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")
HEADING_RE = re.compile(r"^(?P<marks>#{1,3})\s+(?P<text>.+)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(?P<text>.+)$")
BULLET_RE = re.compile(r"^-\s+(?P<text>.+)$")
NUMBER_PREFIX_RE = re.compile(r"^\d+(?:\.\d+)*\s+")
INLINE_RE = re.compile(r"(`[^`]+`|\*[^*]+\*)")

SVG_NS = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
SVG_EXT_URI = "{96DAC541-7B7A-43D3-8B79-37D633B846F1}"


def normalized(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip()).upper()


def add_inline_runs(paragraph, text: str) -> None:
    """Converte itálico simples e código inline sem depender de um parser amplo."""
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token[1:-1])
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_svg_fallback(doc, picture_run, svg_path: Path) -> None:
    """Acrescenta um SVG vetorial a um desenho que já possui PNG fallback."""
    package = doc.part.package
    partname = package.next_partname("/word/media/image%d.svg")
    svg_part = Part(PackURI(partname), "image/svg+xml", svg_path.read_bytes(), package)
    svg_rid = doc.part.relate_to(svg_part, RT.IMAGE)

    blip = picture_run._r.find(".//" + qn("a:blip"))
    if blip is None:
        raise ValueError(f"Blip PNG não localizado para {svg_path}")
    ext_lst = blip.find("{" + A_NS + "}extLst")
    if ext_lst is None:
        ext_lst = etree.SubElement(blip, "{" + A_NS + "}extLst")
    ext = etree.SubElement(ext_lst, "{" + A_NS + "}ext")
    ext.set("uri", SVG_EXT_URI)
    svg_blip = etree.SubElement(ext, "{" + SVG_NS + "}svgBlip")
    svg_blip.set("{" + R_NS + "}embed", svg_rid)


def add_seq_field(paragraph, label: str, value: int) -> None:
    paragraph.add_run(f"{label} ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {label} \\* ARABIC "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = str(value)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, result, end):
        run = OxmlElement("w:r")
        run.append(element)
        paragraph._p.append(run)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def move_before(paragraph, boundary) -> None:
    boundary.addprevious(paragraph._p)


def remove_chapter_range(doc, start_title: str, end_title: str):
    start = next(
        p for p in doc.paragraphs
        if p.style.name == "Heading 1" and normalized(p.text) == normalized(start_title)
    )
    end = next(
        p for p in doc.paragraphs
        if p.style.name == "Heading 1" and normalized(p.text) == normalized(end_title)
    )
    cursor = start._p
    while cursor is not None and cursor is not end._p:
        following = cursor.getnext()
        cursor.getparent().remove(cursor)
        cursor = following
    return end._p


def parse_blocks(markdown_path: Path):
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    pending_figure = None
    paragraph_lines: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_lines
        if paragraph_lines:
            text = " ".join(item.strip() for item in paragraph_lines)
            paragraph_lines = []
            return ("paragraph", text)
        return None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            block = flush_paragraph()
            if block:
                yield block
            continue
        figure_match = FIGURE_COMMENT_RE.fullmatch(stripped)
        if figure_match:
            block = flush_paragraph()
            if block:
                yield block
            pending_figure = {
                "id": figure_match.group("id").strip(),
                "caption": figure_match.group("caption").strip(),
            }
            continue
        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            block = flush_paragraph()
            if block:
                yield block
            if pending_figure is None:
                raise ValueError(f"Imagem sem comentário FIGURE: {stripped}")
            pending_figure["path"] = image_match.group("path")
            pending_figure["alt"] = image_match.group("alt")
            yield ("figure", pending_figure)
            pending_figure = None
            continue
        heading_match = HEADING_RE.fullmatch(stripped)
        if heading_match:
            block = flush_paragraph()
            if block:
                yield block
            level = len(heading_match.group("marks"))
            title = NUMBER_PREFIX_RE.sub("", heading_match.group("text")).strip()
            yield ("heading", level, title)
            continue
        ordered_match = ORDERED_RE.fullmatch(stripped)
        if ordered_match:
            block = flush_paragraph()
            if block:
                yield block
            yield ("ordered", ordered_match.group("text"))
            continue
        bullet_match = BULLET_RE.fullmatch(stripped)
        if bullet_match:
            block = flush_paragraph()
            if block:
                yield block
            yield ("bullet", bullet_match.group("text"))
            continue
        paragraph_lines.append(stripped)

    block = flush_paragraph()
    if block:
        yield block


def sync(input_docx: Path, markdown: Path, output_docx: Path) -> None:
    doc = Document(input_docx)
    for style_name in ("Figure Caption", "Figure Source"):
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    boundary = remove_chapter_range(doc, "METODOLOGIA", "MODELAGEM DA SOLUÇÃO")
    figure_count = sum(1 for p in doc.paragraphs if p.style.name == "Figure Caption")
    bookmark_id = 3000 + figure_count

    for block in parse_blocks(markdown):
        kind = block[0]
        if kind == "heading":
            _, level, title = block
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(paragraph, title)
            move_before(paragraph, boundary)
        elif kind == "paragraph":
            style = "Figure Source" if block[1].startswith("Fonte:") else "Normal"
            paragraph = doc.add_paragraph(style=style)
            if style == "Figure Source":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline_runs(paragraph, block[1])
            move_before(paragraph, boundary)
        elif kind in {"bullet", "ordered"}:
            style = "List Bullet" if kind == "bullet" else "List Number"
            paragraph = doc.add_paragraph(style=style)
            add_inline_runs(paragraph, block[1])
            move_before(paragraph, boundary)
        elif kind == "figure":
            info = block[1]
            svg_path = (markdown.parent / info["path"]).resolve()
            png_path = svg_path.with_suffix(".png")
            if not svg_path.exists() or not png_path.exists():
                raise FileNotFoundError(f"Figura incompleta: {svg_path} / {png_path}")

            figure = doc.add_paragraph()
            figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
            figure.paragraph_format.first_line_indent = Cm(0)
            figure.paragraph_format.keep_together = True
            figure.paragraph_format.keep_with_next = True
            width = Cm(16.0 if "etapas-metodologicas" in svg_path.name else 15.5)
            picture_run = figure.add_run()
            picture_run.add_picture(str(png_path), width=width)
            add_svg_fallback(doc, picture_run, svg_path)
            doc_pr = picture_run._r.find(".//" + qn("wp:docPr"))
            if doc_pr is not None:
                doc_pr.set("name", info["id"])
                doc_pr.set("descr", info["caption"])
            move_before(figure, boundary)

            figure_count += 1
            caption = doc.add_paragraph(style="Figure Caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.first_line_indent = Cm(0)
            caption.paragraph_format.keep_with_next = True
            add_seq_field(caption, "Figura", figure_count)
            caption.add_run(" – ")
            add_inline_runs(caption, info["caption"])
            bookmark = "_TCC_F_" + re.sub(r"[^A-Za-z0-9_]", "_", info["id"])
            add_bookmark(caption, bookmark, bookmark_id)
            bookmark_id += 1
            move_before(caption, boundary)
        else:
            raise AssertionError(block)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()
    sync(args.input_docx, args.markdown, args.output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
